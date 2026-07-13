"""File parsers for extracting raw text from PDF and DOCX files."""

import io
import logging
from pathlib import Path

from .exceptions import ParsingError, UnsupportedFileTypeError

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF, with pdfplumber fallback."""
    text_parts: list[str] = []

    try:
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text.strip())
        doc.close()
        text = "\n\n".join(text_parts).strip()

        if text:
            return text
        logger.info("PyMuPDF returned no text, trying pdfplumber...")
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed: %s, trying pdfplumber...", exc)

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
        text = "\n\n".join(text_parts).strip()
        if text:
            return text
    except Exception as exc:
        logger.error("pdfplumber extraction failed: %s", exc)

    raise ParsingError("Failed to extract text from PDF", "Both PyMuPDF and pdfplumber returned no text")


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        text_parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        text = "\n\n".join(text_parts).strip()
        if text:
            return text
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)

    raise ParsingError("Failed to extract text from DOCX", str(exc))


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a file based on its extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext == ".docx":
        return parse_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {ext}",
            "Only .pdf and .docx files are supported",
        )
